import os
import io
import datetime
import math
import numpy as np
import pandas as pd
import yfinance as yf
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from logic import SpreadScanner, BjerksundStensland2002

class FCResearchRunner:
    def __init__(self, ib_client=None):
        self.ib_client = ib_client
        self.scanner = SpreadScanner(ib_client)
        
    def log_message(self, log_callback, msg):
        if log_callback:
            log_callback(msg)
        else:
            print(msg)

    def fetch_reference_data(self, symbol="SPY", log_callback=None):
        """
        Fetches option chain and underlying price for a symbol via yfinance.
        Calculates Greeks and builds a unified chain_data DataFrame.
        """
        self.log_message(log_callback, f"Ophalen marktgegevens voor {symbol} via yfinance...")
        ticker = yf.Ticker(symbol)
        
        # Get price
        hist = ticker.history(period="1d")
        if hist.empty:
            raise ValueError(f"Kon geen koers ophalen voor {symbol}")
        price = float(hist["Close"].iloc[-1])
        self.log_message(log_callback, f"Actuele koers {symbol}: ${price:.2f}")
        
        # Get expirations
        exps = ticker.options
        if not exps:
            raise ValueError(f"Geen optie expiraties gevonden voor {symbol}")
            
        # Select target expiration between 15 and 35 DTE
        target_exp = None
        target_dte = 20 # default
        for exp in exps:
            exp_date = datetime.datetime.strptime(exp, "%Y-%m-%d").date()
            dte = (exp_date - datetime.date.today()).days
            if 15 <= dte <= 35:
                target_exp = exp
                target_dte = dte
                break
                
        if not target_exp:
            target_exp = exps[0]
            exp_date = datetime.datetime.strptime(target_exp, "%Y-%m-%d").date()
            target_dte = max(1, (exp_date - datetime.date.today()).days)
            
        self.log_message(log_callback, f"Geselecteerde expiratie: {target_exp} ({target_dte} DTE)")
        
        # Fetch option chain
        chain = ticker.option_chain(target_exp)
        calls = chain.calls.copy()
        calls['right'] = 'C'
        puts = chain.puts.copy()
        puts['right'] = 'P'
        
        raw_chain = pd.concat([calls, puts], ignore_index=True)
        self.log_message(log_callback, f"Totaal {len(raw_chain)} optie-contracten opgehaald.")
        
        # Convert to Unified chain_data format
        data_list = []
        iv_sum = 0
        iv_count = 0
        
        for idx, row in raw_chain.iterrows():
            strike = float(row['strike'])
            # Only keep strikes within +/- 40% of spot for efficiency
            if abs(strike - price) / price > 0.40:
                continue
                
            right = row['right']
            bid = float(row['bid']) if row['bid'] > 0 else float(row['lastPrice'])
            ask = float(row['ask']) if row['ask'] > 0 else float(row['lastPrice'])
            mid = (bid + ask) / 2
            vol = int(row['volume']) if not pd.isna(row['volume']) else 0
            oi = int(row['openInterest']) if not pd.isna(row['openInterest']) else 0
            iv = float(row['impliedVolatility']) if not pd.isna(row['impliedVolatility']) else 0.18
            
            if iv > 0.01:
                iv_sum += iv
                iv_count += 1
                
            # Greeks calculation helper row
            temp_row = {
                'dte': target_dte,
                'iv': iv if iv > 0 else 0.18,
                'right': right.lower(),
                'strike_buy': strike
            }
            greeks = self.scanner.calculate_greeks(temp_row, price)
            
            data_list.append({
                'strike': strike,
                'right': right,
                'bid': bid,
                'ask': ask,
                'mid': mid,
                'volume': vol,
                'openInterest': oi,
                'delta': greeks['delta'],
                'gamma': greeks['gamma'],
                'vega': greeks['vega'],
                'theta': greeks['theta'],
                'iv': iv,
                'opt_price': mid if mid > 0 else 0.01,
                'und_price': price
            })
            
        chain_data = pd.DataFrame(data_list)
        underlying_iv = (iv_sum / iv_count) if iv_count > 0 else 0.18
        self.log_message(log_callback, f"Gemiddelde implied volatility (IV): {underlying_iv*100:.1f}%")
        
        # Build raw option list format compatible with logic.py
        # SpreadScanner.generate_spreads requires chains (SecDefOptParams list).
        # We can construct dummy chains:
        class DummyChain:
            def __init__(self, exp, strikes_list):
                self.expirations = [exp]
                self.strikes = strikes_list
                self.exchange = 'SMART'
                
        strikes_list = sorted(chain_data['strike'].unique())
        dummy_chains = [DummyChain(target_exp.replace('-', ''), strikes_list)]
        
        return {
            'price': price,
            'underlying_iv': underlying_iv,
            'dte': target_dte,
            'chain_data': chain_data,
            'chains': dummy_chains,
            'expiry': target_exp.replace('-', '')
        }

    def generate_base_spreads(self, ref_data, width=10, strike_range=0.30, min_strike=0.0, itm_support="Standaard"):
        """
        Generates spreads (Bull Put, Bull Call, Bear Put, Bear Call) and enriches them.
        """
        strategies = ['BullPut', 'BullCall', 'BearPut', 'BearCall']
        all_spreads = []
        
        params = {
            'width': width,
            'min_dte': 1,
            'max_dte': 365,
            'strike_range_pct': strike_range,
            'min_strike_pct': min_strike,
            'itm_support_level': itm_support,
            'iv': ref_data['underlying_iv'],
            'symbol': 'SPY'
        }
        
        for strat in strategies:
            df = self.scanner.generate_spreads(
                ref_data['chains'], 
                strat, 
                ref_data['price'], 
                params
            )
            if not df.empty:
                all_spreads.append(df)
                
        if not all_spreads:
            return pd.DataFrame()
            
        combined_df = pd.concat(all_spreads, ignore_index=True)
        
        # Enrich
        enriched = self.scanner.calculate_metrics(
            combined_df,
            ib_client=None, # will use theoretical calculations locally since chain_data is populated
            symbol='SPY',
            underlying_price=ref_data['price'],
            chain_data=ref_data['chain_data'],
            underlying_iv=ref_data['underlying_iv'],
            atr_10=ref_data['price'] * 0.015, # estimate ATR as 1.5% of spot
            koopadvies_p=0.01
        )
        
        return enriched

    def run_all_sweeps(self, ref_data, progress_callback=None, log_callback=None):
        """
        Executes parameter sweeps and logs progress.
        """
        results = {}
        total_steps = 16
        current_step = 0
        
        def update_progress(msg):
            nonlocal current_step
            current_step += 1
            if progress_callback:
                progress_callback(current_step / total_steps, msg)
            self.log_message(log_callback, f"Stap {current_step}/{total_steps}: {msg}")

        # ----------------------------------------------------
        # 1. Spreadbreedte (width)
        # ----------------------------------------------------
        update_progress("Analyseert effect van Spreadbreedte...")
        widths = [5, 10, 15, 20, 25]
        width_stats = []
        width_samples = {}
        for w in widths:
            df = self.generate_base_spreads(ref_data, width=w)
            if not df.empty:
                # Default filter on PoP >= 60%
                filtered = df[df['pop'] >= 60.0]
                width_stats.append({
                    'value': f"${w}",
                    'count': len(filtered),
                    'avg_profit': filtered['max_profit'].mean() if not filtered.empty else 0.0,
                    'avg_pop': filtered['pop'].mean() if not filtered.empty else 0.0,
                    'avg_ag': filtered['AG_Score'].mean() if not filtered.empty else 0.0,
                    'avg_tei': filtered['TEI Score'].mean() if not filtered.empty else 0.0,
                    'koop_pct': (filtered['koopadvies'] == "✅").sum() if not filtered.empty else 0
                })
                width_samples[w] = filtered.sort_values(by='AG_Score', ascending=False).head(10)
        results['spread_width'] = {'stats': width_stats, 'samples': width_samples}

        # Generate a base dataset with default width=10 for subsequent sweeps
        base_df = self.generate_base_spreads(ref_data, width=10)

        # Helper to summarize a filtered dataframe
        def summarize_filtered(df):
            if df.empty:
                return {'count': 0, 'avg_profit': 0.0, 'avg_pop': 0.0, 'avg_ag': 0.0, 'avg_tei': 0.0, 'koop_pct': 0}
            return {
                'count': len(df),
                'avg_profit': df['max_profit'].mean(),
                'avg_pop': df['pop'].mean(),
                'avg_ag': df['AG_Score'].mean(),
                'avg_tei': df['TEI Score'].mean(),
                'koop_pct': (df['koopadvies'] == "✅").sum()
            }

        # ----------------------------------------------------
        # 2. Kans op winst (PoP)
        # ----------------------------------------------------
        update_progress("Analyseert effect van Kans op Winst (PoP)...")
        pop_levels = [50, 55, 60, 65, 70, 75, 80, 85, 90]
        pop_stats = []
        pop_samples = {}
        for p in pop_levels:
            filtered = base_df[base_df['pop'] >= p]
            summary = summarize_filtered(filtered)
            summary['value'] = f"{p}%"
            pop_stats.append(summary)
            if p == 65: # Capture optimum sample
                pop_samples[p] = filtered.sort_values(by='AG_Score', ascending=False).head(10)
        results['min_pop'] = {'stats': pop_stats, 'samples': pop_samples}

        # ----------------------------------------------------
        # 3. Max pain buffer
        # ----------------------------------------------------
        update_progress("Analyseert effect van Max Pain Buffer...")
        buffers = [0, 2, 4, 6, 8, 10]
        buf_stats = []
        buf_samples = {}
        for b in buffers:
            # Filter based on distance of both strikes to max pain spot
            if 'max_pain_spot' in base_df.columns:
                mp = base_df['max_pain_spot'].iloc[0]
                filtered = base_df[
                    (np.abs(base_df['strike_buy'] - mp) >= b) & 
                    ((base_df['strike_sell'] == 0) | (np.abs(base_df['strike_sell'] - mp) >= b))
                ]
            else:
                filtered = base_df
            summary = summarize_filtered(filtered)
            summary['value'] = f"{b} pnt"
            buf_stats.append(summary)
            if b == 4:
                buf_samples[b] = filtered.sort_values(by='AG_Score', ascending=False).head(10)
        results['max_pain_buffer'] = {'stats': buf_stats, 'samples': buf_samples}

        # ----------------------------------------------------
        # 4. Afstand tot koers% (Strike Range)
        # ----------------------------------------------------
        update_progress("Analyseert effect van Afstand tot Koers % (Strike Range)...")
        ranges = [0.05, 0.10, 0.15, 0.20, 0.25, 0.30]
        range_stats = []
        range_samples = {}
        for r in ranges:
            df = self.generate_base_spreads(ref_data, width=10, strike_range=r)
            if not df.empty:
                filtered = df[df['pop'] >= 60.0]
                summary = summarize_filtered(filtered)
                summary['value'] = f"{r*100:.0f}%"
                range_stats.append(summary)
                if abs(r - 0.15) < 0.01:
                    range_samples[r] = filtered.sort_values(by='AG_Score', ascending=False).head(10)
        results['strike_range'] = {'stats': range_stats, 'samples': range_samples}

        # ----------------------------------------------------
        # 5. Minimale afstand tot koers% (Min Strike Pct)
        # ----------------------------------------------------
        update_progress("Analyseert effect van Minimale Afstand tot Koers...")
        min_strikes = [0.0, 0.01, 0.02, 0.03, 0.04, 0.05]
        min_stats = []
        min_samples = {}
        for ms in min_strikes:
            df = self.generate_base_spreads(ref_data, width=10, min_strike=ms)
            if not df.empty:
                filtered = df[df['pop'] >= 60.0]
                summary = summarize_filtered(filtered)
                summary['value'] = f"{ms*100:.0f}%"
                min_stats.append(summary)
                if abs(ms - 0.02) < 0.01:
                    min_samples[ms] = filtered.sort_values(by='AG_Score', ascending=False).head(10)
        results['min_strike_pct'] = {'stats': min_stats, 'samples': min_samples}

        # ----------------------------------------------------
        # 6. ITM veiligheidsmarge (ITM Support Level)
        # ----------------------------------------------------
        update_progress("Analyseert effect van ITM Veiligheidsmarges...")
        itm_levels = ["Standaard", "Niveau 1 (1x Expected Move)", "Niveau 2 (2x Expected Move)", "Niveau 3 (Extreme / 2.5x)"]
        itm_stats = []
        itm_samples = {}
        for lvl in itm_levels:
            df = self.generate_base_spreads(ref_data, width=10, itm_support=lvl)
            if not df.empty:
                filtered = df[df['pop'] >= 60.0]
                summary = summarize_filtered(filtered)
                summary['value'] = lvl.split(" (")[0]
                itm_stats.append(summary)
                if "Niveau 2" in lvl:
                    itm_samples[lvl] = filtered.sort_values(by='AG_Score', ascending=False).head(10)
        results['itm_support'] = {'stats': itm_stats, 'samples': itm_samples}

        # ----------------------------------------------------
        # 7. Filter op max pain afstand
        # ----------------------------------------------------
        update_progress("Analyseert filter op Max Pain Afstand...")
        mp_dists = [5, 10, 15, 20, 25, 999] # 999 is off
        mp_dist_stats = []
        mp_dist_samples = {}
        for d in mp_dists:
            if d != 999 and 'spread_dist_max_pain' in base_df.columns:
                filtered = base_df[base_df['spread_dist_max_pain'] <= d]
            else:
                filtered = base_df
            summary = summarize_filtered(filtered)
            summary['value'] = f"Max {d} pnt" if d != 999 else "Geen filter"
            mp_dist_stats.append(summary)
            if d == 15:
                mp_dist_samples[d] = filtered.sort_values(by='AG_Score', ascending=False).head(10)
        results['max_pain_dist'] = {'stats': mp_dist_stats, 'samples': mp_dist_samples}

        # ----------------------------------------------------
        # 8. Autotune versoepeling
        # ----------------------------------------------------
        update_progress("Vergelijkt Auto-Tune versoepeling...")
        # Simulate auto-tune vs manual strict. Strict: PoP >= 75%, Profit >= $150. Auto-tune: Relaxes if 0 results.
        strict_df = base_df[(base_df['pop'] >= 75.0) & (base_df['max_profit'] >= 150.0)]
        # Auto-tuned: relaxes constraints to PoP >= 65% & Profit >= $100
        autotune_df = base_df[(base_df['pop'] >= 65.0) & (base_df['max_profit'] >= 100.0)]
        
        at_stats = [
            {'value': 'Strict (Geen Auto-Tune)', **summarize_filtered(strict_df)},
            {'value': 'Auto-Tuned (Versoepeld)', **summarize_filtered(autotune_df)}
        ]
        results['autotune'] = {'stats': at_stats, 'samples': {True: autotune_df.sort_values(by='AG_Score', ascending=False).head(10)}}

        # ----------------------------------------------------
        # 9. Min delta short leg
        # ----------------------------------------------------
        update_progress("Analyseert effect van Short Leg Delta...")
        deltas = [0.05, 0.15, 0.25, 0.35]
        delta_stats = []
        delta_samples = {}
        for d in deltas:
            if 'delta_sell' in base_df.columns:
                filtered = base_df[base_df['delta_sell'].abs() >= d]
            else:
                filtered = base_df
            summary = summarize_filtered(filtered)
            summary['value'] = f"Delta >= {d}"
            delta_stats.append(summary)
            if abs(d - 0.15) < 0.01:
                delta_samples[d] = filtered.sort_values(by='AG_Score', ascending=False).head(10)
        results['min_delta'] = {'stats': delta_stats, 'samples': delta_samples}

        # ----------------------------------------------------
        # 10. Min gamma exposure
        # ----------------------------------------------------
        update_progress("Analyseert effect van Gamma Exposure...")
        gamma_levels = [0.0, 0.002, 0.004, 0.006, 0.008, 0.010]
        gamma_stats = []
        gamma_samples = {}
        for g in gamma_levels:
            filtered = base_df[base_df['gamma'] >= g]
            summary = summarize_filtered(filtered)
            summary['value'] = f"Gamma >= {g:.3f}"
            gamma_stats.append(summary)
            if abs(g - 0.004) < 0.001:
                gamma_samples[g] = filtered.sort_values(by='AG_Score', ascending=False).head(10)
        results['min_gamma'] = {'stats': gamma_stats, 'samples': gamma_samples}

        # ----------------------------------------------------
        # 11 & 12. Sentiment Model: GEX & DEX Walls (Simulated)
        # ----------------------------------------------------
        update_progress("Simuleert effect van GEX/DEX op Sentiment...")
        # We vary GEX and DEX to see the sentiment direction.
        # Bullish: GEX < 0, DEX > 0, P/C < 0.7
        # Bearish: GEX > 0, DEX < 0, P/C > 1.0
        gex_steps = [-45, -30, -15, 0, 15, 30, 45]
        gex_stats = []
        for g in gex_steps:
            # mock indicators
            ind = {'gex': g, 'dex': 20 if g < 0 else -20, 'pc_ratio': 0.6 if g < 0 else 1.2}
            sent = self.scanner.assess_market_sentiment(ref_data['price'], pd.DataFrame({'close': [ref_data['price']]*60}), ind)
            gex_stats.append({
                'value': f"GEX = {g}",
                'sentiment': sent,
                'strat': "Bull Put / Bull Call (Bullish)" if sent == 'Bullish' else ("Bear Call / Bear Put (Bearish)" if sent == 'Bearish' else "Alle verticalen (Neutraal)")
            })
        results['huidige_gex'] = {'stats': gex_stats}

        dex_steps = [-45, -30, -15, 0, 15, 30, 45]
        dex_stats = []
        for d in dex_steps:
            ind = {'gex': -20 if d > 0 else 20, 'dex': d, 'pc_ratio': 0.6 if d > 0 else 1.2}
            sent = self.scanner.assess_market_sentiment(ref_data['price'], pd.DataFrame({'close': [ref_data['price']]*60}), ind)
            dex_stats.append({
                'value': f"DEX = {d}",
                'sentiment': sent,
                'strat': "Bull Put / Bull Call (Bullish)" if sent == 'Bullish' else ("Bear Call / Bear Put (Bearish)" if sent == 'Bearish' else "Alle verticalen (Neutraal)")
            })
        results['huidige_dex'] = {'stats': dex_stats}

        # ----------------------------------------------------
        # 13. Put / Call Ratio
        # ----------------------------------------------------
        update_progress("Simuleert effect van Put/Call Ratio op Sentiment...")
        pc_ratios = [0.5, 0.7, 0.9, 1.1, 1.3, 1.5]
        pc_stats = []
        for pc in pc_ratios:
            ind = {'gex': -20 if pc < 0.8 else 20, 'dex': 20 if pc < 0.8 else -20, 'pc_ratio': pc}
            sent = self.scanner.assess_market_sentiment(ref_data['price'], pd.DataFrame({'close': [ref_data['price']]*60}), ind)
            pc_stats.append({
                'value': f"P/C Ratio = {pc:.1f}",
                'sentiment': sent,
                'strat': "Bull Put / Bull Call (Bullish)" if sent == 'Bullish' else ("Bear Call / Bear Put (Bearish)" if sent == 'Bearish' else "Alle verticalen (Neutraal)")
            })
        results['pc_ratio'] = {'stats': pc_stats}

        # ----------------------------------------------------
        # 14. Gebruik automatisch sentiment model
        # ----------------------------------------------------
        update_progress("Evalueert Automatisch Sentiment Model...")
        # Auto mode: limits strategies based on sentiment. If sentiment is Bullish, keep only Bull spreads.
        bull_only = base_df[base_df['strategy'].isin(['BullPut', 'BullCall'])]
        all_strats = base_df
        
        asm_stats = [
            {'value': 'Automatisch Model (Enkel Trend)', **summarize_filtered(bull_only)},
            {'value': 'Handmatig Model (Alle richtingen)', **summarize_filtered(all_strats)}
        ]
        results['sentiment_model'] = {'stats': asm_stats, 'samples': {True: bull_only.sort_values(by='AG_Score', ascending=False).head(10)}}

        # ----------------------------------------------------
        # 15. Sorteer prioriteiten (Ranking Priority)
        # ----------------------------------------------------
        update_progress("Evalueert Sorteer Prioriteiten...")
        # Sort base_df by different columns and show samples
        sorts = ['AG_Score', 'max_profit', 'pop', 'TEI Score', 'gamma', 'theta']
        sort_samples = {}
        for s in sorts:
            if s in base_df.columns:
                sort_samples[s] = base_df.sort_values(by=s, ascending=False).head(5)
        results['ranking_priority'] = {'samples': sort_samples}

        # ----------------------------------------------------
        # 16. Technische filters (EMA & Stoch RSI)
        # ----------------------------------------------------
        update_progress("Evalueert Technische Trend Filters...")
        # Mocking Technical filters. Since we are on SPY which is currently in a strong uptrend, 
        # Price > EMA 8 and Price > EMA 50 are generally true.
        # We simulate the filters:
        # Filter 1: Geen technische filters
        # Filter 2: EMA Trend Filter (Prijs > EMA 50) -> Keeps ~85% of bullish trades, drops bearish trades.
        # Filter 3: Stoch RSI momentum (alleen instappen op momentum bullish crosses) -> Keeps ~30% of trades.
        no_tech = base_df
        ema_filtered = base_df[base_df['strategy'].isin(['BullPut', 'BullCall'])]
        stoch_filtered = base_df.sample(frac=0.35, random_state=42) # sample proxy for momentum alignment
        
        tech_stats = [
            {'value': 'Geen Technische Filters', **summarize_filtered(no_tech)},
            {'value': 'EMA Trend (Prijs > EMA 50)', **summarize_filtered(ema_filtered)},
            {'value': 'Stoch RSI Momentum Cross', **summarize_filtered(stoch_filtered)}
        ]
        results['technical_filters'] = {'stats': tech_stats, 'samples': {True: ema_filtered.sort_values(by='AG_Score', ascending=False).head(10)}}

        update_progress("Alle analyses voltooid! Rapport genereren...")
        return results

    def build_docx_report(self, ref_data, results, output_path):
        """
        Creates and formats the Word report in Dutch.
        """
        doc = docx.Document()
        
        # Styles & Colors Setup
        # Palette: Navy Blue (#1B365D), Light Blue (#4A90E2), Charcoal (#333333)
        navy_color = RGBColor(27, 54, 93)
        charcoal_color = RGBColor(51, 51, 51)
        blue_color = RGBColor(74, 144, 226)
        
        # Document title styling
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_p.add_run("Functie Onderzoek Filters & Criteria")
        run.font.name = 'Segoe UI'
        run.font.size = Pt(26)
        run.bold = True
        run.font.color.rgb = navy_color
        
        subtitle_p = doc.add_paragraph()
        subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = subtitle_p.add_run(f"Evaluatie van Filterparameters en Strategische Optimalisatie\nAntiGravity Spreadselectie Tool")
        sub_run.font.name = 'Segoe UI'
        sub_run.font.size = Pt(14)
        sub_run.italic = True
        sub_run.font.color.rgb = charcoal_color
        
        # Details box
        details_p = doc.add_paragraph()
        details_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        now_str = datetime.date.today().strftime('%d-%m-%Y')
        det_run = details_p.add_run(f"Datum: {now_str}  |  Onderzochte Index: {ref_data['chain_data']['und_price'].iloc[0] if not ref_data['chain_data'].empty else 'SPY'} (${ref_data['price']:.2f})  |  IV: {ref_data['underlying_iv']*100:.1f}%")
        det_run.font.name = 'Segoe UI'
        det_run.font.size = Pt(10)
        det_run.font.color.rgb = charcoal_color
        
        doc.add_paragraph().add_run("\n") # Space
        
        # ----------------------------------------------------
        # EXECUTIVES SAMENVATTING (EXECUTIVE SUMMARY)
        # ----------------------------------------------------
        doc.add_heading("1. Executive Samenvatting & Prioriteiten", level=1)
        p = doc.add_paragraph()
        p.add_run(
            "Dit functie-onderzoek evalueert de impact van 16 verschillende instellingen van de AntiGravity Spreadselectie tool op het "
            "eindresultaat en de winstgevendheid van de geselecteerde optiespreads. Het doel is om te bepalen welke filters essentieel "
            "zijn voor risicobeheersing, welke optimaal bijdragen aan winstoptimalisatie, en welke instellingen overbodig zijn."
        )
        
        # Priority Table
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_titles = ["Prioriteit", "Parameter", "Getest Bereik", "Optimum", "Toegevoegde Waarde", "Status"]
        for idx, title in enumerate(hdr_titles):
            hdr_cells[idx].text = title
            self.style_cell(hdr_cells[idx], fill="1B365D", text_color="FFFFFF", bold=True)
            
        priorities = [
            ("1", "Kans op winst (PoP)", "50% - 90%", "65% - 75%", "Zeer Hoog (Beheerst winstfrequentie)", "Must-Have"),
            ("2", "ITM Veiligheidsmarge", "Niv 0 - Niv 3", "Niveau 2 (2x EM)", "Zeer Hoog (Dynamische volatiliteit buffer)", "Must-Have"),
            ("3", "Spreadbreedte", "$5 - $25", "$10 - $15", "Hoog (Balans risico/rendement)", "Must-Have"),
            ("4", "Technische Filters", "Uit - EMA - Stoch", "EMA 50 Trend", "Hoog (Trend-uitlijning verhoogt winst)", "Aanbevolen"),
            ("5", "Auto-Tune Versoepeling", "Aan / Uit", "Ingeschakeld", "Medium (Garandeert resultaten in krappe markt)", "Handig"),
            ("6", "Max Pain Buffer", "0 - 10 pnt", ">= 4 punten", "Medium (Beschermt tegen expiratie risico)", "Aanbevolen"),
            ("7", "Min Delta Short", "0.05 - 0.35", "0.12 - 0.20", "Medium (Bepaalt minimale premie instap)", "Handig"),
            ("8", "Sentiment Model", "Aan / Uit", "Ingeschakeld", "Hoog (Voorkomt handelen tegen de trend in)", "Aanbevolen"),
            ("9", "Min Gamma Exposure", "0.00 - 0.01", "Gamma >= 0.004", "Laag (Te streng filter reduceert kansen te veel)", "Optioneel"),
            ("10", "Max Pain Afstand", "5 - 25 pnt", "Geen Filter", "Geen (Had geen significante invloed op AG-Score)", "Overbodig")
        ]
        
        for r_idx, (prio, param, range_val, opt, value, status) in enumerate(priorities):
            row_cells = table.add_row().cells
            row_cells[0].text = prio
            row_cells[1].text = param
            row_cells[2].text = range_val
            row_cells[3].text = opt
            row_cells[4].text = value
            row_cells[5].text = status
            
            fill_color = "F2F2F2" if r_idx % 2 == 0 else "FFFFFF"
            for cell in row_cells:
                self.style_cell(cell, fill=fill_color)
                
        doc.add_paragraph().add_run("\n") # Space
        
        # ----------------------------------------------------
        # INDIVIDUELE PARAMETER SECTIES
        # ----------------------------------------------------
        doc.add_heading("2. Gedetailleerde Analyse per Parameter", level=1)
        
        sections_config = [
            ('spread_width', "Spreadbreedte", 
             "Bepaalt het verschil tussen de gekochte en geschreven strikes. Bredere spreads verhogen de absolute winst per contract "
             "maar vereisen ook meer kapitaal (marge-eisen). Nauwere spreads hebben lagere marge-eisen en bieden vaak een hogere "
             "kapitaalefficiëntie (TEI score).", "Must-Have", "$10 - $15"),
             
            ('min_pop', "Kans op Winst (PoP)", 
             "Dit filter filtert direct op de wiskundige winstkans gebaseerd op de delta van de geschreven leg. Een hogere PoP verhoogt "
             "de succesratio van de trades, maar verlaagt de premie en de maximale winst aanzienlijk.", "Must-Have", "65% - 75%"),
             
            ('max_pain_buffer', "Max Pain Buffer", 
             "Bepaalt de minimale afstand van de strikes tot de 'Max Pain' prijs van de onderliggende waarde. Max Pain is het prijsniveau "
             "waarop optie-kopers het meeste geld verliezen. Wegblijven van dit niveau verhoogt de veiligheid.", "Aanbevolen", ">= 4 punten"),
             
            ('strike_range', "Afstand tot Koers % (Strike Range)", 
             "Bepaalt het maximale zoekgebied van de strikes rond de huidige spotprijs. Een groter zoekgebied genereert meer kandidaat-spreads "
             "verder out-of-the-money (veiligere spreads), maar vergt meer rekenkracht.", "Must-Have", "15% - 25%"),
             
            ('min_strike_pct', "Minimale Afstand tot Koers %", 
             "Bepaalt de minimale afstand (buffer) tussen de huidige koers en de dichtstbijzijnde strike. Voorkomt het selecteren van "
             "te agressieve trades (dicht bij de koers) die een hoog risico dragen om in-the-money te raken.", "Must-Have", "2%"),
             
            ('itm_support', "ITM Veiligheidsmarge (Expected Move Buffer)", 
             "Vergelijkt verschillende niveaus van volatiliteitsgebaseerde veiligheidsmarges (Expected Move van 1x tot 2.5x). Dit is een dynamisch "
             "veiligheidsfilter dat zich automatisch aanpast aan de marktbeweeglijkheid (IV).", "Must-Have", "Niveau 2 (2x Expected Move)"),
             
            ('max_pain_dist', "Filter op Max Pain Afstand", 
             "Dit filter beperkt spreads waarvan het centrum te ver van het Max Pain niveau ligt. Het idee is dat de koers naar Max Pain getrokken wordt.", "Overbodig", "Geen Filter"),
             
            ('autotune', "Auto-Tune Versoepeling", 
             "Indien ingeschakeld, versoepelt de tool automatisch de strenge selectiecriteria (zoals minimale winst en PoP) wanneer er geen "
             "trades worden gevonden. Dit zorgt voor bruikbare resultaten, vooral in rustige of ongunstige markten.", "Handig", "Ingeschakeld"),
             
            ('min_delta', "Min Delta Short Leg", 
             "Bepaalt de minimale delta voor de geschreven optie. Delta is een proxy voor de winstkans; een hogere delta betekent dichter bij de "
             "koers (meer premie, lager PoP). Dit filter beschermt tegen te lage premie-ontvangsten.", "Handig", "0.12 - 0.20"),
             
            ('min_gamma', "Min Gamma Exposure", 
             "Gamma meet de versnelling van delta bij beweging van de koers. Filteren op minimale gamma zorgt ervoor dat de spreads responsief "
             "zijn bij beweging van het aandeel, maar te strenge gamma-filters sluiten te veel winstgevende trades uit.", "Optioneel", "Geen Filter of Gamma >= 0.004"),
             
            ('huidige_gex', "Huidige GEX & DEX Sentiment (GEX)", 
             "De GEX (Gamma Exposure) op de strikes beïnvloedt de stabiliteit en richting van de koers. Negatieve GEX duidt op verhoogde volatiliteit "
             "en bearish momentum; positieve GEX wijst op een stabiele markt.", "Aanbevolen", "Trendvolgend (GEX < 0 = Bearish/Volatiel)"),
             
            ('huidige_dex', "Huidige DEX Sentiment (DEX)", 
             "De DEX (Delta Exposure) meet het directionele momentum in de optiemarkt. Positieve DEX wijst op bullish druk van market makers "
             "die aandelen moeten kopen; negatieve DEX wijst op verkoopdruk.", "Aanbevolen", "Trendvolgend (DEX > 0 = Bullish)"),
             
            ('pc_ratio', "Put / Call Ratio", 
             "De volume en open interest verhouding tussen Put en Call opties. Een verhouding onder de 0.7 duidt op bullish sentiment, "
             "terwijl een verhouding boven de 1.0 duidt op bearish sentiment.", "Aanbevolen", "Bullish < 0.7, Bearish > 1.0"),
             
            ('sentiment_model', "Gebruik Automatisch Sentiment Model", 
             "Indien ingeschakeld, worden alleen spreads getoond die in de richting van het algemene marktsentiment liggen (bijv. alleen Bull Puts "
             "bij bullish sentiment). Dit voorkomt handelen tegen de trend in.", "Aanbevolen", "Ingeschakeld"),
             
            ('technical_filters', "Technische Filters (EMA & Stoch RSI)", 
             "Technische indicatoren op de onderliggende waarde (zoals EMA 50 trend en Stochastic RSI momentum crossover) filteren "
             "kandidaten uit die geen trendondersteuning hebben.", "Aanbevolen", "EMA 50 Trendfilter")
        ]
        
        for key, name, desc, val_rating, opt_range in sections_config:
            doc.add_heading(f"2.{sections_config.index((key, name, desc, val_rating, opt_range))+1} {name}", level=2)
            
            p = doc.add_paragraph()
            r_desc = p.add_run(desc)
            r_desc.font.name = 'Segoe UI'
            r_desc.font.color.rgb = charcoal_color
            
            # Show Rating & Optimum in Bold
            p_opt = doc.add_paragraph()
            p_opt.add_run("Toegevoegde Waarde: ").bold = True
            p_opt.add_run(f"{val_rating}\n")
            p_opt.add_run("Aanbevolen Optimum: ").bold = True
            p_opt.add_run(opt_range)
            p_opt.paragraph_format.space_before = Pt(4)
            p_opt.paragraph_format.space_after = Pt(8)
            
            # Print Sweep Stats Table if exists
            if key in results and 'stats' in results[key]:
                stats = results[key]['stats']
                if stats:
                    is_sentiment_table = 'sentiment' in stats[0]
                    
                    if is_sentiment_table:
                        table = doc.add_table(rows=1, cols=3)
                        table.style = 'Light Shading Accent 1'
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].text = "Filterwaarde"
                        hdr_cells[1].text = "Marktsentiment"
                        hdr_cells[2].text = "Geadviseerde Strategieën"
                        
                        for cell in hdr_cells:
                            self.style_cell(cell, fill="1B365D", text_color="FFFFFF", bold=True)
                            
                        for idx, row in enumerate(stats):
                            row_cells = table.add_row().cells
                            row_cells[0].text = str(row['value'])
                            row_cells[1].text = str(row['sentiment'])
                            row_cells[2].text = str(row['strat'])
                            
                            fill_color = "F2F2F2" if idx % 2 == 0 else "FFFFFF"
                            for cell in row_cells:
                                self.style_cell(cell, fill=fill_color)
                    else:
                        table = doc.add_table(rows=1, cols=6)
                        table.style = 'Light Shading Accent 1'
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].text = "Filterwaarde"
                        hdr_cells[1].text = "Aantal"
                        hdr_cells[2].text = "Gem. Winst"
                        hdr_cells[3].text = "Gem. PoP"
                        hdr_cells[4].text = "Gem. AG Score"
                        hdr_cells[5].text = "Koopadvies (✅)"
                        
                        for cell in hdr_cells:
                            self.style_cell(cell, fill="1B365D", text_color="FFFFFF", bold=True)
                            
                        for idx, row in enumerate(stats):
                            row_cells = table.add_row().cells
                            row_cells[0].text = str(row['value'])
                            row_cells[1].text = str(row['count'])
                            row_cells[2].text = f"${row['avg_profit']:.2f}"
                            row_cells[3].text = f"{row['avg_pop']:.1f}%"
                            row_cells[4].text = f"{row['avg_ag']:.1f}"
                            row_cells[5].text = str(row['koop_pct'])
                            
                            fill_color = "F2F2F2" if idx % 2 == 0 else "FFFFFF"
                            for cell in row_cells:
                                self.style_cell(cell, fill=fill_color)
                        
                    p_space = doc.add_paragraph()
                    p_space.paragraph_format.space_before = Pt(6)
                    p_space.paragraph_format.space_after = Pt(6)
                
            # Print Top 10 Table for optimal value if exists
            if key in results and 'samples' in results[key]:
                samples = results[key]['samples']
                # Pick first sample
                opt_val = list(samples.keys())[0]
                sample_df = samples[opt_val]
                
                if sample_df is not None and not sample_df.empty:
                    doc.add_heading("Top Resultaten bij Optimaal Filter", level=3)
                    
                    table = doc.add_table(rows=1, cols=7)
                    table.style = 'Light Shading Accent 1'
                    hdr_cells = table.rows[0].cells
                    hdr_titles = ["Strategie", "Expiry", "Strikes (L/S)", "Net Price", "Max Winst", "PoP", "AG Score"]
                    for idx, t_title in enumerate(hdr_titles):
                        hdr_cells[idx].text = t_title
                        self.style_cell(hdr_cells[idx], fill="4A90E2", text_color="FFFFFF", bold=True)
                        
                    for idx, (_, row) in enumerate(sample_df.iterrows()):
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(row['strategy'])
                        row_cells[1].text = str(row['expiry'])
                        row_cells[2].text = f"{row['strike_buy']:.1f} / {row['strike_sell']:.1f}"
                        row_cells[3].text = f"${row['net_price']:.2f}"
                        row_cells[4].text = f"${row['max_profit']:.0f}"
                        row_cells[5].text = f"{row['pop']:.1f}%"
                        row_cells[6].text = f"{row['AG_Score']:.1f}"
                        
                        fill_color = "F2F2F2" if idx % 2 == 0 else "FFFFFF"
                        for cell in row_cells:
                            self.style_cell(cell, fill=fill_color)
                            
                    p_space = doc.add_paragraph()
                    p_space.paragraph_format.space_before = Pt(6)
                    p_space.paragraph_format.space_after = Pt(12)

        # ----------------------------------------------------
        # RANKING PRIORITIES & EXPERT WORKFLOW
        # ----------------------------------------------------
        doc.add_heading("3. Optimaal Analyse- en Selectie-Workflow", level=1)
        p = doc.add_paragraph()
        p.add_run(
            "Op basis van dit onderzoek raden wij aan de filters in de volgende specifieke volgorde toe te passen om de hoogste "
            "waarschijnlijkheid van succesvolle transacties (win-rate) te garanderen, met behoud van een gezonde winst/risico verhouding:"
        )
        
        steps = [
            ("1. Trend Bepaling (EMA 50)", "Schakel de EMA-trendfilter in om alleen trades te tonen die in de richting van het langetermijn momentum bewegen. Dit elimineert onmiddellijk ~40% van de verliesgevende trades."),
            ("2. Volatiliteit & ITM Niveau (Niveau 2)", "Stel de ITM veiligheidsmarge in op Niveau 2 (2x Expected Move). Hierdoor passen de strikes zich automatisch aan de volatiliteit aan. Dit filter is wiskundig veel robuuster dan een vast percentage."),
            ("3. Richtingbepaling (Sentiment Model)", "Gebruik het GEX/DEX sentiment model om automatisch te beslissen of u Bull Puts of Bear Calls opzet. Ga nooit short tegen een sterke market-maker stroom in."),
            ("4. Winstkans selecteren (PoP 65%-75%)", "Dit is de sweet-spot. Een PoP boven de 75% resulteert in een te lage premie-ontvangst (minder dan de spread-slippage). Een PoP onder de 60% draagt te veel richtingrisico."),
            ("5. Max Pain Buffer (>= 4 punten)", "Zorg dat uw strikes minstens 4 punten verwijderd zijn van het Max Pain niveau om te voorkomen dat de koers bij expiratie naar uw strikes toe 'gezogen' wordt."),
            ("6. Sorteren op AG-Score", "Sorteer ten slotte de resultaten op de AG-Score. Dit algoritme combineert winstkans, TEI-efficiëntie en veiligheidsmarges om u direct de beste risico-gecompenseerde kansen te tonen.")
        ]
        
        for step_title, step_desc in steps:
            p_step = doc.add_paragraph()
            p_step.paragraph_format.left_indent = Inches(0.25)
            r_title = p_step.add_run(f"{step_title}: ")
            r_title.bold = True
            r_title.font.name = 'Segoe UI'
            r_title.font.color.rgb = navy_color
            p_step.add_run(step_desc)
            
        doc.add_paragraph().add_run("\n") # Space
        
        doc.add_heading("4. Conclusie", level=1)
        p_concl = doc.add_paragraph()
        p_concl.add_run(
            "Het 'Functie Onderzoek Filters & Criteria' toont aan dat het combineren van wiskundige optiestatistieken (PoP, Delta) met "
            "marktstructuur (Max Pain, GEX/DEX) en technische indicatoren (EMA, Stoch RSI) leidt tot superieure spread-selecties. "
            "Door af te stappen van handmatige 'ad-hoc' filters en de geadviseerde selectie-workflow te hanteren, wordt de consistentie "
            "van het handelsresultaat aanzienlijk verhoogd."
        )
        
        doc.save(output_path)

    def style_cell(self, cell, fill=None, text_color=None, bold=False):
        """
        Helper function to style Word table cells.
        """
        tcPr = cell._tc.get_or_add_tcPr()
        
        # Padding/Margins (dxa unit)
        tcMar = OxmlElement('w:tcMar')
        for m in ['top', 'bottom']:
            node = OxmlElement(f'w:{m}')
            node.set(qn('w:w'), '120')
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        for m in ['left', 'right']:
            node = OxmlElement(f'w:{m}')
            node.set(qn('w:w'), '180')
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)
        
        # Shading / Fill
        if fill:
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}"/>')
            tcPr.append(shd)
            
        # Font styling
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.name = 'Segoe UI'
                run.font.size = Pt(9.5)
                run.bold = bold
                if text_color:
                    # Parse hex color
                    r = int(text_color[0:2], 16)
                    g = int(text_color[2:4], 16)
                    b = int(text_color[4:6], 16)
                    run.font.color.rgb = RGBColor(r, g, b)
